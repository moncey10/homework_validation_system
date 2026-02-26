# app.py
import os
import io
import re
import json
from typing import Dict, Any, List, Tuple

import requests
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from PIL import Image, ImageOps, ImageFilter
import pytesseract
import os

# Serve static files from outputs directory
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from dotenv import load_dotenv
load_dotenv()

# Optional extractors for DOCX/PDF
try:
    from docx import Document  # python-docx
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    import reportlab
except Exception as e:
    reportlab = None
    print(f"[WARN] reportlab import failed: {e}")

try:
    from pdf2image import convert_from_bytes  # requires poppler
except Exception:
    convert_from_bytes = None

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ✅ Gemini SDK
try:
    from google import genai
except Exception as e:
    genai = None
    print(f"[WARN] google-genai import failed: {e}")

# ✅ Google Cloud Vision SDK (for better handwritten OCR)
try:
    from google.cloud import vision
    from google.cloud.vision_v1 import types
    google_vision_available = True
except Exception as e:
    google_vision_available = False
    print(f"[WARN] google-cloud-vision import failed: {e}")



app = FastAPI()
app.mount("/files", StaticFiles(directory="outputs"), name="files")


outputs_dir = "outputs"
os.makedirs(outputs_dir, exist_ok=True)

app.mount("/outputs", StaticFiles(directory=outputs_dir), name="outputs")

# Create outputs directory if it doesn't exist
outputs_dir = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(outputs_dir, exist_ok=True)

@app.get("/outputs/{filename}")
async def get_output_file(filename: str):
    """Serve files from the outputs directory."""
    filepath = os.path.join(outputs_dir, filename)
    if os.path.exists(filepath):
        return FileResponse(filepath)
    raise HTTPException(status_code=404, detail="File not found")

@app.get("/storage/{filename}")
async def get_storsge_file(filename:str):
    """Serve files from the storage directory."""
    filepath = os.path.join(outputs_dir, filename)
    if os.path.exists(filepath):
        return FileResponse(filepath)
    raise HTTPException(status_code=404, detail="File not found")

@app.get("/debug/erp-row")
async def debug_erp_row(homework_id: int, student_id: int):
    """Debug endpoint: shows the raw ERP row so you can see all field names."""
    try:
        row = fetch_student_record(homework_id, student_id)
        return {"erp_row": row, "keys": list(row.keys())}
    except Exception as e:
        return {"error": str(e)}

@app.get("/debug/env")
def debug_env():
    return {
        "has_gemini_keys": bool(GOOGLE_API_KEYS),
        "num_keys": len(GOOGLE_API_KEYS),
        "has_openai_key": bool(os.getenv("OPENAI_API_KEY")),
    }
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



if os.name == "nt":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"



ERP_BASE = os.getenv("ERP_BASE", "https://erp.triz.co.in/lms_data")
STORAGE_BASE = os.getenv("STORAGE_BASE", "https://erp.triz.co.in/storage/student/")
ERP_TOKEN = os.getenv("ERP_TOKEN", "")


def get_public_base_url() -> str:
    """
    Returns the public base URL of this server.
    Priority:
      1. SPACE_HOST  — set automatically by Hugging Face Spaces (most reliable)
      2. HF_SPACE    — manual fallback env var for HF
      3. APP_BASE_URL — custom deployment domain
      4. localhost   — local dev only
    """
    hf_host = os.getenv("SPACE_HOST", "").strip()
    if hf_host:
        return f"https://{hf_host}"

    hf_space = os.getenv("HF_SPACE", "").strip()
    if hf_space:
        return f"https://{hf_space}"

    custom = os.getenv("APP_BASE_URL", "").strip()
    if custom:
        return custom.rstrip("/")

    return "http://127.0.0.1:7860"


def build_pdf_url(filename: str) -> str:
    """Given a saved PDF filename, return its full public URL."""
    if not filename:
        return ""
    return f"{get_public_base_url()}/outputs/{filename}"


def make_question_marks(mcq_results: list) -> list:
    """
    Convert internal mcq_results into a clean list the frontend can use
    to show ✓ ✗ ○ next to each question number.

    Each item:
      {
        "qid":            "Q1",
        "mark":           "correct" | "wrong" | "unattempted",
        "student_answer": "A",          # what the student chose (empty if unattempted)
        "correct_answer": "B"           # the right answer (null if unknown)
      }
    """
    result = []
    for r in (mcq_results or []):
        if r.get('unattempted'):
            mark = "unattempted"
        elif r.get('correct') is True:
            mark = "correct"
        else:
            mark = "wrong"
        result.append({
            "qid":            r.get('qid', ''),
            "mark":           mark,
            "student_answer": r.get('chosen', ''),
            "correct_answer": r.get('correct_answer'),
        })
    return result


# API Key Rotation - Support multiple API keys for higher limits
GOOGLE_API_KEYS = []
for i in range(1, 10):  # Support up to 10 API keys
    key = os.getenv(f"GOOGLE_API_KEY_{i}", "").strip()
    if key:
        GOOGLE_API_KEYS.append(key)

# Fallback to single key if no multi-key config
GOOGLE_API_KEY = (os.getenv("GOOGLE_API_KEY") or "").strip()
if not GOOGLE_API_KEYS and GOOGLE_API_KEY:
    GOOGLE_API_KEYS = [GOOGLE_API_KEY]

# Track current key index and rate-limited keys
current_key_index = 0
rate_limited_keys = set()  # Track keys that are rate limited

GEMINI_MODEL = (os.getenv("GEMINI_MODEL", "models/gemini-flash-lite-latest") or "").strip()
if GEMINI_MODEL and not GEMINI_MODEL.startswith("models/"):
    GEMINI_MODEL = "models/" + GEMINI_MODEL


GOOGLE_CLOUD_VISION_API_KEY = (os.getenv("GCV_API_KEY") or "").strip()
# Fall back to Gemini API key if no separate Vision key provided
if not GOOGLE_CLOUD_VISION_API_KEY and GOOGLE_API_KEY:
    GOOGLE_CLOUD_VISION_API_KEY = GOOGLE_API_KEY

vision_client = None
if google_vision_available and GOOGLE_CLOUD_VISION_API_KEY:
    try:
        # Use API key authentication
        vision_client = vision.ImageAnnotatorClient(client_options={
            'api_key': GOOGLE_CLOUD_VISION_API_KEY
        })
        print("[INFO] Google Cloud Vision client initialized")
    except Exception as e:
        print(f"[WARN] Google Cloud Vision init failed: {e}")

gemini_client = None
GEMINI_LAST_ERROR = ""

def _init_gemini_client(key_index: int = 0) -> None:
    """Initialize Gemini client with the API key at the given index."""
    global gemini_client, GEMINI_LAST_ERROR, current_key_index
    
    current_key_index = key_index
    
    if not GOOGLE_API_KEYS:
        GEMINI_LAST_ERROR = "No GOOGLE_API_KEY configured"
        gemini_client = None
        return
    
    if key_index >= len(GOOGLE_API_KEYS):
        GEMINI_LAST_ERROR = "All API keys rate limited or exhausted"
        gemini_client = None
        return
    
    api_key = GOOGLE_API_KEYS[key_index]

    if not genai:
        GEMINI_LAST_ERROR = "google-genai not installed / import failed"
        gemini_client = None
        return

    if not api_key:
        GEMINI_LAST_ERROR = f"GOOGLE_API_KEY_{key_index + 1} not set"
        gemini_client = None
        return

    try:
        gemini_client = genai.Client(api_key=api_key)
        GEMINI_LAST_ERROR = ""
        print(f"[INFO] Gemini client initialized with key #{key_index + 1}")
    except Exception as e:
        gemini_client = None
        GEMINI_LAST_ERROR = str(e)
        print(f"[WARN] Gemini init failed: {GEMINI_LAST_ERROR}")


def _is_rate_limit_error(error_msg: str) -> bool:
    """Check if the error is a rate limit error (429) or service unavailable (503)."""
    if not error_msg:
        return False
    lower = error_msg.lower()
    return ("429" in lower or 
            "503" in lower or 
            "rate_limit" in lower or 
            "resource_exhausted" in lower or 
            "rate limit" in lower or
            "unavailable" in lower)


def _rotate_to_next_key() -> bool:
    """Rotate to the next available API key. Returns True if successful, False if all keys exhausted."""
    global current_key_index, rate_limited_keys
    
    if len(GOOGLE_API_KEYS) <= 1:
        return False
    
    # Mark current key as rate limited
    rate_limited_keys.add(current_key_index)
    print(f"[WARN] Key #{current_key_index + 1} rate limited, rotating to next key...")
    
    # Find next available key
    attempts = 0
    while attempts < len(GOOGLE_API_KEYS):
        current_key_index = (current_key_index + 1) % len(GOOGLE_API_KEYS)
        if current_key_index not in rate_limited_keys:
            _init_gemini_client(current_key_index)
            if gemini_client:
                print(f"[INFO] Rotated to API key #{current_key_index + 1}")
                return True
        attempts += 1
    
    # All keys exhausted
    GEMINI_LAST_ERROR = "All API keys are rate limited"
    print("[ERROR] All API keys are rate limited")
    return False


_init_gemini_client(0)


def parse_gemini_error(error_msg: str) -> dict:
    msg = (error_msg or "").strip()
    lower = msg.lower()

    if "service_disabled" in lower or "generativelanguage.googleapis.com" in lower:
        return {"ok": False, "error_type": "GEMINI_SERVICE_DISABLED", "message": msg}

    if "api key" in lower or "invalid" in lower or "permission" in lower or "unauthorized" in lower:
        return {"ok": False, "error_type": "GEMINI_KEY_OR_PERMISSION_ERROR", "message": msg}

    return {"ok": False, "error_type": "GEMINI_ERROR", "message": msg}



def extract_qid_from_prompt(prompt: str, erp_row: dict = None) -> str:
    """
    Extract the question number (e.g. 'Q5') from the ERP row or prompt string.
    Priority:
      1. Direct field in erp_row: question_no, q_no, sr_no, serial_no, qno, question_number, q_number, order, position
      2. Pattern match in prompt text: 'Q5:', 'Question 5:', '5.', '5)'
      3. Falls back to 'Q1' if nothing found.
    """
    import re as _re

    # Priority 1: check ERP row fields directly
    if erp_row and isinstance(erp_row, dict):
        for field in ("question_no", "q_no", "qno", "sr_no", "serial_no",
                      "question_number", "q_number", "order", "position",
                      "question_order", "q_order", "seq", "sequence", "index"):
            val = erp_row.get(field)
            if val is not None:
                try:
                    num = int(str(val).strip())
                    if 1 <= num <= 200:
                        print(f"[INFO] extract_qid: found Q{num} from erp_row['{field}']={val}")
                        return f"Q{num}"
                except (ValueError, TypeError):
                    pass

    # Priority 2: parse from prompt text
    p = (prompt or "").strip()
    m = _re.match(r'^[Qq]\s*(\d+)', p)
    if m:
        return f"Q{m.group(1)}"
    m2 = _re.match(r'^Question\s*(\d+)', p, _re.IGNORECASE)
    if m2:
        return f"Q{m2.group(1)}"
    m3 = _re.match(r'^(\d+)[.)\s]', p)
    if m3:
        return f"Q{m3.group(1)}"
    first_line = p.split('\n')[0]
    m4 = _re.search(r'[Qq]\s*(\d+)', first_line)
    if m4:
        return f"Q{m4.group(1)}"

    print(f"[WARN] extract_qid: could not determine question number from prompt={repr(p[:80])} erp_row_keys={list((erp_row or {}).keys())}")
    return "Q1"


def generate_gemini_response(
    prompt: str,
    system_prompt: str = "",
    max_tokens: int = 650,
    temperature: float = 0.3,
) -> str:
    global GEMINI_LAST_ERROR, gemini_client, rate_limited_keys

    if gemini_client is None:
        if not GEMINI_LAST_ERROR:
            GEMINI_LAST_ERROR = "Gemini client not initialized"
        # Try to reinitialize if we have keys available
        if GOOGLE_API_KEYS and current_key_index not in rate_limited_keys:
            _init_gemini_client(current_key_index)
        if gemini_client is None:
            return ""

    try:
        contents = []
        if system_prompt:
            contents.append(system_prompt)
        contents.append(prompt)

        resp = gemini_client.models.generate_content(
            model=GEMINI_MODEL,
            contents=contents,
            config={"temperature": temperature, "max_output_tokens": max_tokens},
        )
        text = (getattr(resp, "text", "") or "").strip()
        if text:
            GEMINI_LAST_ERROR = ""
        return text
    except Exception as e:
        error_msg = str(e)
        print(f"[ERROR] Gemini call failed: {error_msg}")
        
        # Check if it's a rate limit error and try to rotate
        if _is_rate_limit_error(error_msg):
            GEMINI_LAST_ERROR = error_msg
            if _rotate_to_next_key():
                # Retry with new key
                return generate_gemini_response(prompt, system_prompt, max_tokens, temperature)
        
        GEMINI_LAST_ERROR = error_msg
        return ""

import time

def generate_gemini_with_retry(prompt: str, system_prompt: str, max_tokens=450, temperature=0.3, retries=3) -> str:
    last = ""
    for i in range(retries):
        text = generate_gemini_response(
            prompt=prompt,
            system_prompt=system_prompt,
            max_tokens=max_tokens,
            temperature=temperature,
        )
        if text:
            return text
        last = GEMINI_LAST_ERROR
        # small backoff
        time.sleep(1 + i)
    return ""

def cheap_overlap_score(student_text: str, prompt: str) -> int:
    # remove tiny words
    def tokens(s):
        return {w for w in re.findall(r"[a-zA-Z]{4,}", (s or "").lower())}
    s = tokens(student_text)
    p = tokens(prompt)
    if not s or not p:
        return 0
    overlap = len(s & p) / max(1, len(p))
    # map to a sane range
    return int(round(min(0.6, overlap) * 100))  # cap at 60



def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def cosine_sim(a: str, b: str) -> float:
    a = (a or "").strip().lower()
    b = (b or "").strip().lower()
    if not a or not b:
        return 0.0
    vec = TfidfVectorizer().fit([a, b])
    X = vec.transform([a, b])
    return float(cosine_similarity(X[0], X[1])[0][0])


def normalize_level(level: str) -> str:
    l = (level or "").strip().lower()
    if l in ("easy",):
        return "Easy"
    if l in ("hard",):
        return "Hard"
    if l in ("meadium", "mediam", "medium"):
        return "Medium"
    return "Medium"


def level_policy(student_level: str) -> dict:
    lvl = normalize_level(student_level).lower()
    if lvl == "easy":
        return {"w_sim": 0.8, "w_cov": 0.2, "verified": 65, "partial": 40, "kp_thr": 0.25}
    if lvl == "hard":
        return {"w_sim": 0.4, "w_cov": 0.6, "verified": 85, "partial": 65, "kp_thr": 0.40}
    return {"w_sim": 0.6, "w_cov": 0.4, "verified": 75, "partial": 55, "kp_thr": 0.20}


def mcq_partial_credit(student_level: str) -> dict:
    """
    Returns partial credit percentage for MCQ questions based on student level.
    This allows easier students to get partial marks even if they get some questions wrong.
    
    Returns dict with:
    - credit_per_question: percentage earned per correct answer
    - passing_threshold: minimum percentage needed to pass
    """
    lvl = normalize_level(student_level).lower()
    if lvl == "easy":
        # Easy students get 50% credit per correct answer
        return {"credit_per_question": 50, "passing_threshold": 50}
    if lvl == "hard":
        # Hard students need 100% - no partial credit
        return {"credit_per_question": 100, "passing_threshold": 100}
    # Medium students get 75% credit per correct answer
    return {"credit_per_question": 75, "passing_threshold": 75}


def keypoint_coverage(student_text: str, key_points: List[str], kp_threshold: float) -> Tuple[List[str], List[str], float]:
    covered, missing = [], []
    for kp in key_points:
        kp = (kp or "").strip()
        if not kp:
            continue
        s = cosine_sim(kp, student_text)
        if s >= kp_threshold:
            covered.append(kp)
        else:
            missing.append(kp)

    total = len(covered) + len(missing)
    coverage = (len(covered) / total) if total else 0.0
    return covered, missing, coverage



def infer_question_type_from_prompt(prompt: str, student_text: str = "") -> str:
    p = _norm(prompt)

    # Explicit markers - check for (mcq) first since it's common in parentheses
    if re.search(r"\(mcq\)", p) or re.search(r"\btype\s*:\s*mcq\b", p) or re.search(r"\bquestion_type\s*:\s*mcq\b", p):
        return "mcq"
    if re.search(r"\btype\s*:\s*narrative\b", p) or re.search(r"\bquestion_type\s*:\s*narrative\b", p):
        return "narrative"

    # Heuristic: options A/B/C/D exist in prompt -> likely MCQ
    if re.search(r"\b(a|b|c|d)\s*[\)\.]\s+", p) or "option a" in p or "option b" in p:
        return "mcq"
    
    # Check if prompt contains common MCQ keywords
    if re.search(r"\bchoose the correct|which is correct|select the right|multiple choice|single answer\b", p):
        return "mcq"

    # Check student answer for MCQ indicators if provided
    if student_text:
        s = _norm(student_text)
        # If student answer contains Option A/B/C/D, treat as MCQ
        if re.search(r"\boption\s*[a-d]\b", s) or re.search(r"^\(?\s*[a-d]\s*\)?$", s.strip()):
            return "mcq"
        # If answer starts with A. or B. etc.
        if re.search(r"^[a-d]\.\s+", s.strip()):
            return "mcq"

    return "narrative"


def parse_questions_from_prompt(prompt: str) -> List[Dict[str, Any]]:
    """
    Parse individual questions from the prompt, detecting MCQ vs Narrative for each.
    Returns list of dicts with: qid, type, question_text, correct_answer (for MCQ)
    """
    questions = []
    # Match patterns like "Q1:", "Q2.", "Question 1:", etc.
    q_pattern = re.compile(r'(Q\s*\d+[.:]\s*|Question\s*\d+[.:]\s*)(.*?)(?=(Q\s*\d|Question\s*\d|$))', re.IGNORECASE | re.DOTALL)
    
    # Alternative: split by Q1, Q2, etc.
    lines = prompt.split('\n')
    current_q = None
    current_type = None
    current_qid = None
    current_correct = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect new question
        q_match = re.match(r'^(Q\s*\d+|Question\s*\d+)[.:]\s*(.*)', line, re.IGNORECASE)
        if q_match:
            # Save previous question if exists
            if current_q is not None:
                questions.append({
                    'qid': current_qid,
                    'type': current_type,
                    'question': current_q,
                    'correct_answer': current_correct
                })
            # Start new question
            current_qid = q_match.group(1).strip()
            remaining = q_match.group(2).strip()
            current_q = remaining
            current_type = None
            current_correct = None
            
            # Check if this is MCQ or Narrative
            line_lower = line.lower()
            if '(mcq)' in line_lower or 'multiple choice' in line_lower or 'type: mcq' in line_lower:
                current_type = 'mcq'
            elif 'narrative' in line_lower or 'type: narrative' in line_lower:
                current_type = 'narrative'
        else:
            # This line belongs to current question
            if current_q is not None:
                current_q += ' ' + line
                
                # Check for type markers
                line_lower = line.lower()
                if current_type is None:
                    if '(mcq)' in line_lower or 'multiple choice' in line_lower or 'type: mcq' in line_lower:
                        current_type = 'mcq'
                    elif 'narrative' in line_lower or 'type: narrative' in line_lower:
                        current_type = 'narrative'
                
                # Check for correct answer (for MCQ)
                if current_type == 'mcq':
                    # First check: is this line "Correct Answer(s):" with nothing after it?
                    # If so, we need to look for the answer on the next line
                    if re.search(r'^correct\s*answer\s*\(?s\)?\s*[:\.]?\s*$', line, re.IGNORECASE):
                        # Set flag to look for answer on next line
                        current_q += ' [CORRECT_ANSWER_PENDING]'
                        continue
                    
                    # Check if we have a pending correct answer marker
                    if '[CORRECT_ANSWER_PENDING]' in current_q:
                        # This line should contain the answer like "A. Devdatta"
                        letter_match = re.search(r'^([A-D])\.?\s*', line)
                        if letter_match:
                            current_correct = letter_match.group(1).upper()
                            # Remove the pending marker from question text
                            current_q = current_q.replace(' [CORRECT_ANSWER_PENDING]', '')
                            continue
                        else:
                            # Not a letter, remove the pending marker
                            current_q = current_q.replace(' [CORRECT_ANSWER_PENDING]', '')
                    
                    # Look for "Correct Answer(s):" or "Correct:" or "Answer:" in same line
                    # Support formats: "Correct Answer(s): A.", "Correct: B", "Answer: C"
                    correct_match = re.search(r'(?:Correct\s*(?:Answer)?|Answer)[:.]\s*(?:[A-D]\.?\s*)?(.+)', line, re.IGNORECASE)
                    if correct_match and not current_correct:
                        # Extract just the letter (A, B, C, or D)
                        correct_text = correct_match.group(1).strip()
                        letter_match = re.search(r'^([A-D])\b', correct_text)
                        if letter_match:
                            current_correct = letter_match.group(1).upper()
                        else:
                            # Try to extract first letter
                            current_correct = correct_text[0].upper() if correct_text else None
    
    # Don't forget the last question
    if current_q is not None:
        questions.append({
            'qid': current_qid,
            'type': current_type,
            'question': current_q,
            'correct_answer': current_correct
        })
    
    # If no questions parsed, fall back to old behavior
    if not questions:
        qtype = infer_question_type_from_prompt(prompt)
        return [{'qid': extract_qid_from_prompt(prompt), 'type': qtype, 'question': prompt, 'correct_answer': None}]
    
    return questions


def extract_mcq_choice(text: str) -> str:
    """
    Extract chosen option from student text:
    supports: A, (B), Option C, Ans: D, Answer: B
    """
    t = _norm(text)

    m = re.search(r"\b(answer|ans|selected)\s*[:\-]?\s*\(?\s*([a-d])\s*\)?\b", t)
    if m:
        return m.group(2)

    m2 = re.search(r"\boption\s*([a-d])\b", t)
    if m2:
        return m2.group(1)

    m3 = re.search(r"^\(?\s*([a-d])\s*\)?$", t.strip())
    if m3:
        return m3.group(1)

    # last-resort: find first standalone A/B/C/D
    m4 = re.search(r"\b([a-d])\b", t)
    if m4:
        return m4.group(1)

    return ""


def extract_mcq_answers_with_qid(text: str) -> Dict[str, str]:
    """
    Extract MCQ answers WITH question numbers from student text.
    This handles shuffled answers where question numbers are needed to match.
    
    Supports patterns like:
    - "Q1: A, Q2: C, Q3: B"
    - "Q1. A Q2. C Q3. B"
    - "1) A 2) C 3) B"
    - "Answer 1: A Answer 2: C Answer 3: B"
    - "Q1 A Q2 C Q3 B" (space separated)
    
    Returns dict like: {"Q1": "A", "Q2": "C", "Q3": "B"}
    """
    results = {}
    t = (text or "").strip()
    
    if not t:
        return results

    # Pattern 1: Q1: A, Q2. B, Q3 - C, Question 4: D
    pattern1 = re.compile(r'(Q(?:uestion)?\s*(\d+))[:.\-\s]+([a-dA-D])', re.IGNORECASE)
    for match in pattern1.finditer(t):
        qnum = match.group(2)
        answer = match.group(3).upper()
        results[f"Q{qnum}"] = answer
    
    # Pattern 2: 1) A, 2) B, 3: C (numbered without Q prefix)
    pattern2 = re.compile(r'(?:^|\s)(\d+)\s*[\):\.]\s*([a-dA-D])(?:\s|$)', re.IGNORECASE)
    for match in pattern2.finditer(t):
        qnum = match.group(1)
        answer = match.group(2).upper()
        # Only add if not already found (Q pattern takes priority)
        if f"Q{qnum}" not in results:
            results[f"Q{qnum}"] = answer
    
    # Pattern 3: "Answer for Q1 is A", "Answer to question 2: B"
    pattern3 = re.compile(r'(?:answer|ans)\s*(?:for|to)?\s*(?:Q(?:uestion)?\s*)?(\d+)\s*(?:is|was)?\s*[:\-]?\s*([a-dA-D])', re.IGNORECASE)
    for match in pattern3.finditer(t):
        qnum = match.group(1)
        answer = match.group(2).upper()
        if f"Q{qnum}" not in results:
            results[f"Q{qnum}"] = answer
    
    # Pattern 4: Line by line format like "Q1 A" or "1 A" on same line
    pattern4 = re.compile(r'(?:^|\n)\s*(Q(?:uestion)?\s*)?(\d+)\s+([a-dA-D])\s*(?:\n|\s{2,}|$)', re.IGNORECASE)
    for match in pattern4.finditer(t):
        qnum = match.group(2)
        answer = match.group(3).upper()
        if f"Q{qnum}" not in results:
            results[f"Q{qnum}"] = answer
    
    return results


def extract_correct_mcq_from_prompt(prompt: str) -> str:
    """
    This is IMPORTANT:
    Your prompt must contain correct option somewhere like:
      - Correct: B
      - Answer: C
      - correct_option: D
      - Correct Answer(s): A. Devdatta
    or JSON: {"correct_option":"B"}
    
    Supports formats:
      - "Correct Answer: A"
      - "Correct Answer(s): A. Devdatta"
      - "Correct: B"
      - "Answer: C"
    """
    p = (prompt or "").strip()
    if not p:
        return ""

    # JSON prompt support
    if p.startswith("{") and p.endswith("}"):
        try:
            obj = json.loads(p)
            for k in ("correct_option", "correct", "answer", "ans"):
                v = obj.get(k)
                if isinstance(v, str) and v.strip():
                    return extract_mcq_choice(v)
        except Exception:
            pass

    # Text prompt support - new format: "Correct Answer(s): A. Devdatta" or "Correct Answer: B"
    t = _norm(p)
    
    # Pattern 1: "Correct Answer(s): A. ..." or "Correct Answer: B. ..."
    # This handles format like "Correct Answer(s): A. Devdatta" or "Correct Answer(s):
    #    A. Devdatta"
    m1 = re.search(r"correct\s*answer\s*\(?s\)?\s*[:\.]\s*([a-d])\.?\s*", t)
    if m1:
        return m1.group(1)
    
    # Pattern 1b: Handle multi-line format where answer is on next line like:
    # "Correct Answer(s):\n   A. Devdatta"
    m1b = re.search(r"correct\s*answer\s*\(?s\)?\s*[:\.]\s*\n\s*([a-d])\.?", t)
    if m1b:
        return m1b.group(1)
    
    # Pattern 1c: Handle format with option text after letter like "Correct Answer(s): A. Devdatta"
    m1c = re.search(r"correct\s*answer\s*\(?s\)?\s*[:\.]\s*([a-d])\.", t)
    if m1c:
        return m1c.group(1)
    
    # Pattern 2: "Correct: A" or "Answer: B" (original pattern)
    m = re.search(r"\b(correct|answer|ans)\s*[:\-]?\s*\(?\s*([a-d])\s*\)?\b", t)
    if m:
        return m.group(2)

    return ""



def _erp_get(params: dict) -> list:
    headers = {}
    if ERP_TOKEN:
        headers["Authorization"] = f"Bearer {ERP_TOKEN}"

    r = requests.get(ERP_BASE, params=params, headers=headers, timeout=30)
    r.raise_for_status()
    data = r.json()
    if not isinstance(data, list):
        raise HTTPException(status_code=502, detail="ERP returned invalid JSON (expected list).")
    return data


def fetch_student_record(homework_id: int, student_id: int) -> Dict[str, Any]:
    data = _erp_get({"table": "homework", "filters[id]": homework_id, "filters[student_id]": student_id})
    if not data:
        raise HTTPException(status_code=404, detail="No ERP record found for this homework_id + student_id")
    return data[0]


def fetch_student_level_from_erp(row: Dict[str, Any]) -> str:
    """
    ERP field name is not guaranteed; try common ones.
    """
    for k in ("student_level", "level", "difficulty", "difficulty_level"):
        v = row.get(k)
        if isinstance(v, str) and v.strip():
            return normalize_level(v)
    return "Medium"



def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """
    Enhanced preprocessing for better OCR on handwritten images.
    Includes adaptive thresholding, noise removal, and contrast enhancement.
    """
    # Convert to grayscale
    img = img.convert("L")
    
    w, h = img.size
    
    # Scale up for better detail (especially for handwritten)
    if max(w, h) < 2000:
        scale = 2000 / max(w, h)
        new_w = int(w * scale)
        new_h = int(h * scale)
        img = img.resize((new_w, new_h), Image.LANCZOS)
    
    # Apply adaptive thresholding for better handwritten recognition
    from PIL import ImageFilter
    
    # Try multiple preprocessing approaches and use the best
    img_enhanced = img
    
    # Method 1: Increase contrast significantly
    img_contrast = img.point(lambda p: 255 if p > 180 else int(p * 1.5))
    
    # Method 2: Apply sharpening twice for handwritten
    img_sharp = img.filter(ImageFilter.SHARPEN)
    img_sharp = img_sharp.filter(ImageFilter.SHARPEN)
    
    # Method 3: Apply unsharp mask for edge enhancement
    img_unsharp = img.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))
    
    # Use the sharpened version as primary
    img = img_sharp
    
    # Apply binary threshold with lower cutoff to capture lighter handwriting
    img = img.point(lambda p: 255 if p > 160 else 0)
    
    return img


def _extract_text_google_vision(image_bytes: bytes) -> str:
    """
    Extract text using Google Cloud Vision API - much better for handwriting.
    Returns empty string if API is not available.
    """
    global vision_client
    
    if not vision_client:
        return ""
    
    try:
        # Create image object
        image = vision.Image(content=image_bytes)
        
        # Use document text detection for better handwriting
        response = vision_client.document_text_detection(image=image)
        
        if response.texts:
            return "\n".join([t.description for t in response.texts])
        return ""
    except Exception as e:
        print(f"[WARN] Google Vision OCR failed: {e}")
        return ""


def extract_text_from_image(image_bytes: bytes, filename: str = "unknown") -> str:
    if not image_bytes or len(image_bytes) < 50:
        raise HTTPException(status_code=400, detail=f"Invalid file: '{filename}' - empty/too small")

    valid_image_signatures = {
        b"\xff\xd8\xff": "JPEG",
        b"\x89PNG\r\n\x1a\n": "PNG",
        b"GIF87a": "GIF",
        b"GIF89a": "GIF",
        b"BM": "BMP",
    }
    is_valid = any(image_bytes.startswith(sig) for sig in valid_image_signatures)
    if not is_valid:
        head = image_bytes[:12]
        raise HTTPException(status_code=400, detail=f"Invalid image format: '{filename}' (header={head})")

    # First try Google Cloud Vision (better for handwriting)
    if vision_client:
        gv_text = _extract_text_google_vision(image_bytes)
        if gv_text and len(gv_text.strip()) > 10:
            return _clean_extracted_text(gv_text)
    
    # Fallback to Tesseract with improved preprocessing
    try:
        img = Image.open(io.BytesIO(image_bytes))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid image '{filename}': {e}")

    img = _preprocess_for_ocr(img)

    # Try multiple OCR configurations for better handwritten recognition
    ocr_configs = [
        "--oem 3 --psm 6",  # Default
        "--oem 3 --psm 4",  # Treat as single column
        "--oem 1 --psm 3",  # Fully automatic
    ]
    
    best_text = ""
    best_confidence = 0
    
    for config in ocr_configs:
        try:
            text = pytesseract.image_to_string(img, lang="eng", config=config)
            if text and len(text.strip()) > len(best_text.strip()):
                best_text = text
        except Exception:
            continue
    
    if not best_text:
        # Fallback to default if all fail
        try:
            best_text = pytesseract.image_to_string(img, lang="eng", config="--oem 3 --psm 6")
        except pytesseract.TesseractNotFoundError:
            raise HTTPException(status_code=500, detail="Tesseract OCR not found. Install it / fix path.")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"OCR failed: {e}")

    text = (best_text or "").strip()
    text = re.sub(r"[ \t]+", " ", text)
    return text


def _clean_extracted_text(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_docx(docx_bytes: bytes, filename: str = "unknown.docx") -> str:
    if Document is None:
        raise HTTPException(status_code=500, detail="DOCX support not installed. Add 'python-docx'.")
    try:
        doc = Document(io.BytesIO(docx_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _clean_extracted_text("\n".join(parts))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Unable to read DOCX '{filename}': {e}")


def extract_text_from_pdf(pdf_bytes: bytes, filename: str = "unknown.pdf") -> Dict[str, Any]:
    used_ocr = False
    extracted = ""

    if PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            extracted = _clean_extracted_text("\n\n".join(parts))
        except Exception:
            extracted = ""

    if len(extracted) < 50:
        if convert_from_bytes is None:
            return {"text": extracted, "used_ocr": False, "needs_ocr": True}
        try:
            used_ocr = True
            # Higher DPI for better handwritten OCR
            pages = convert_from_bytes(pdf_bytes, dpi=300)
            page_texts = []
            for img in pages:
                # Use the improved preprocessing
                img = _preprocess_for_ocr(img)
                
                # Try multiple OCR configs
                for config in ["--oem 3 --psm 6", "--oem 3 --psm 4", "--oem 1 --psm 3"]:
                    try:
                        t = pytesseract.image_to_string(img, lang="eng", config=config) or ""
                        if t.strip() and len(t.strip()) > 20:
                            page_texts.append(t)
                            break
                    except:
                        continue
            
            if page_texts:
                extracted = _clean_extracted_text("\n\n".join(page_texts))
            else:
                # Final fallback with default config
                img = pages[0] if pages else None
                if img:
                    img = _preprocess_for_ocr(img)
                    extracted = pytesseract.image_to_string(img, lang="eng", config="--oem 3 --psm 6") or ""
        except Exception as e:
            return {"text": extracted, "used_ocr": used_ocr, "needs_ocr": True, "ocr_error": str(e)}

    return {"text": extracted, "used_ocr": used_ocr, "needs_ocr": False}


def get_question_positions_from_pdf(pdf_bytes: bytes) -> Dict[int, List[Dict]]:
    """
    Detect question number positions in a PDF.
    Strategy 1: pypdf text-layer visitor (fast, for PDFs with text layer).
    Strategy 2: pdf2image + pytesseract OCR (for image-based PDFs).
    Returns dict: page_num -> [{qid, y_pos, x_pos}] in PDF coords (origin bottom-left).
    """
    try:
        from pypdf import PdfReader
        from io import BytesIO
        reader = PdfReader(BytesIO(pdf_bytes))
        question_positions: Dict[int, List[Dict]] = {}

        def _normalise_ocr_qid(token: str):
            t = token.strip().rstrip('.')
            m = re.match(r'^[Qq]\s*(\d+)$', t)
            if m:
                return f"Q{m.group(1)}"
            ocr_map = {'i': '1', 'I': '1', 'l': '1', 'o': '0', 'O': '0',
                       'z': '2', 'Z': '2', 's': '5', 'S': '5', 'g': '9'}
            m2 = re.match(r'^[Qq]([a-zA-Z\d])$', t)
            if m2:
                digit = ocr_map.get(m2.group(1), m2.group(1))
                if digit.isdigit():
                    return f"Q{digit}"
            return None

        for page_num, page in enumerate(reader.pages):
            page_height = float(page.mediabox.height) if hasattr(page.mediabox, 'height') else 792
            page_width  = float(page.mediabox.width)  if hasattr(page.mediabox, 'width')  else 595
            found: List[Dict] = []
            existing_qids: set = set()

            # Strategy 1: text layer
            try:
                parts = []
                def _visitor(text, cm, tm, font_dict, font_size):
                    if text and text.strip():
                        parts.append((text.strip(), float(tm[4]) if tm else 0, float(tm[5]) if tm else 0))
                page.extract_text(visitor_text=_visitor)
                tl_patterns = [
                    re.compile(r'\bQ\s*(\d+)\b', re.IGNORECASE),
                    re.compile(r'\bQuestion\s*(\d+)\b', re.IGNORECASE),
                    re.compile(r'^(\d+)[.):\s]'),
                ]
                for text_frag, x, y in parts:
                    for pat in tl_patterns:
                        m = pat.match(text_frag)
                        if m:
                            qid = f"Q{m.group(1)}"
                            if qid not in existing_qids:
                                existing_qids.add(qid)
                                found.append({'qid': qid, 'y_pos': y, 'x_pos': x})
                            break
            except Exception as tl_err:
                print(f"[WARN] text-layer page {page_num}: {tl_err}")

            # Strategy 2: OCR fallback (image-based PDFs)
            if not found:
                try:
                    from pdf2image import convert_from_bytes as _c2b
                    import pytesseract
                    rendered = _c2b(pdf_bytes, dpi=72, first_page=page_num+1, last_page=page_num+1)
                    if rendered:
                        img = rendered[0]
                        img_w, img_h = img.size
                        scale_x = page_width  / img_w
                        scale_y = page_height / img_h
                        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                        for i, token in enumerate(ocr_data['text']):
                            if not token or not token.strip() or int(ocr_data['conf'][i]) < 20:
                                continue
                            dm = re.match(r'^[Qq]\s*(\d+)[.:]?$', token.strip())
                            qid = f"Q{dm.group(1)}" if dm else _normalise_ocr_qid(token)
                            if qid and qid not in existing_qids:
                                pdf_x = ocr_data['left'][i] * scale_x
                                pdf_y = page_height - (ocr_data['top'][i] + ocr_data['height'][i] * 0.5) * scale_y
                                existing_qids.add(qid)
                                found.append({'qid': qid, 'y_pos': pdf_y, 'x_pos': pdf_x})
                except Exception as ocr_err:
                    print(f"[WARN] OCR fallback page {page_num}: {ocr_err}")

            if found:
                found.sort(key=lambda d: -d['y_pos'])
                question_positions[page_num] = found
        return question_positions
    except Exception as e:
        print(f"[WARN] Failed to get question positions: {e}")
        return {}


def create_annotated_pdf(
    original_pdf_bytes: bytes,
    mcq_results: List[Dict[str, Any]] = None,
    match_percentage: int = 0,
    status: str = "Needs Review",
    student_level: str = "Medium",
    question_type: str = "mcq"
) -> bytes:
    """
    Annotate every question number found in the PDF with a coloured mark:
      Correct     -> filled green circle + white tick  (✓)
      Wrong       -> filled red circle   + white cross (✗)
      Unattempted -> hollow orange circle              (○)

    Any question detected in the PDF that has NO entry in mcq_results is
    automatically treated as unattempted (hollow orange circle).
    """
    if not reportlab:
        print("[WARN] reportlab not available, returning original PDF")
        return original_pdf_bytes

    try:
        from pypdf import PdfWriter, PdfReader
        from io import BytesIO

        question_positions = get_question_positions_from_pdf(original_pdf_bytes)
        print(f"[INFO] Detected question positions: {question_positions}")

        qid_location: Dict[str, tuple] = {}
        for pg, items in question_positions.items():
            for item in items:
                qid_location[item["qid"]] = (pg, item["y_pos"], item["x_pos"])

        results_by_qid: Dict[str, Dict] = {}
        for r in (mcq_results or []):
            qid = r.get("qid", "")
            if qid:
                results_by_qid[qid] = r

        def _draw_mark(c, x, y, is_correct, is_unattempted, radius=14):
            if is_correct and not is_unattempted:
                # ✓ Green filled circle — correct only
                c.setStrokeColor(colors.Color(0.0, 0.65, 0.0))
                c.setFillColor(colors.Color(0.0, 0.65, 0.0))
                c.setLineWidth(2)
                c.circle(x, y, radius, fill=1)
                c.setFillColor(colors.white)
                c.setFont("Helvetica-Bold", int(radius * 1.5))
                c.drawString(x - radius * 0.5, y - radius * 0.45, "\u2713")
            else:
                # ✗ Red filled circle — wrong OR unattempted
                c.setStrokeColor(colors.Color(0.85, 0.1, 0.1))
                c.setFillColor(colors.Color(0.85, 0.1, 0.1))
                c.setLineWidth(2)
                c.circle(x, y, radius, fill=1)
                c.setFillColor(colors.white)
                c.setFont("Helvetica-Bold", int(radius * 1.5))
                c.drawString(x - radius * 0.5, y - radius * 0.45, "\u2717")

        MARK_RADIUS   = 14
        MARK_X_OFFSET = -(MARK_RADIUS + 4)

        original_reader = PdfReader(BytesIO(original_pdf_bytes))
        writer = PdfWriter()
        total_pages = len(original_reader.pages)

        for page_num, page in enumerate(original_reader.pages):
            page_width  = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            packet = BytesIO()
            c = canvas.Canvas(packet, pagesize=(page_width, page_height))

            # Draw a mark for every detected question on this page
            for item in question_positions.get(page_num, []):
                qid   = item["qid"]
                y_pos = item["y_pos"]
                x_pos = item["x_pos"]

                result         = results_by_qid.get(qid)
                is_unattempted = True   # default: no data → unattempted
                is_correct     = False

                if result is not None:
                    explicit_unattempted = result.get("unattempted")
                    chosen = result.get("chosen", "")
                    correct_val = result.get("correct")

                    if explicit_unattempted is True:
                        # Explicitly flagged as unattempted
                        is_unattempted = True
                        is_correct = False
                    elif not chosen or str(chosen).strip() == "":
                        # No answer recorded → treat as unattempted
                        is_unattempted = True
                        is_correct = False
                    else:
                        # Answer was given — mark correct or wrong
                        is_unattempted = False
                        is_correct = bool(correct_val)

                mark_x = max(MARK_RADIUS + 2, x_pos + MARK_X_OFFSET)
                mark_y = y_pos + MARK_RADIUS * 0.3
                _draw_mark(c, mark_x, mark_y, is_correct, is_unattempted, MARK_RADIUS)

            # Fallback: results whose qid was not detected in the PDF
            undetected = [r for r in (mcq_results or []) if r.get("qid") not in qid_location]
            if undetected:
                per_page  = max(1, (len(undetected) + total_pages - 1) // total_pages)
                start_idx = page_num * per_page
                page_slice = undetected[start_idx: start_idx + per_page]
                y_start   = page_height - 100
                y_spacing = max(20, (page_height - 130) / max(1, per_page))
                for i, result in enumerate(page_slice):
                    explicit_unattempted = result.get("unattempted")
                    chosen = result.get("chosen", "")
                    correct_val = result.get("correct")

                    if explicit_unattempted is True or not chosen or str(chosen).strip() == "":
                        is_unattempted = True
                        is_correct = False
                    else:
                        is_unattempted = False
                        is_correct = bool(correct_val)
                    y_pos = y_start - i * y_spacing
                    if y_pos < 30:
                        break
                    _draw_mark(c, 18, y_pos, is_correct, is_unattempted, 9)

            # Header bar on first page
            if page_num == 0:
                header_h = 58
                c.setFillColor(colors.Color(0.93, 0.93, 0.93))
                c.rect(0, page_height - header_h, page_width, header_h, fill=1, stroke=0)

                sc = (colors.Color(0.0, 0.65, 0.0) if status == "Verified"
                      else colors.Color(1.0, 0.55, 0.0) if status == "Partial"
                      else colors.Color(0.85, 0.1, 0.1))

                c.setFillColor(sc); c.setStrokeColor(sc)
                c.circle(18, page_height - 22, 8, fill=1)
                c.setFont("Helvetica-Bold", 14)
                c.drawString(34, page_height - 27, f"Status: {status}")

                c.setFillColor(colors.black)
                c.setFont("Helvetica-Bold", 14)
                c.drawString(page_width * 0.42, page_height - 27, f"Score: {match_percentage}%")
                c.drawString(page_width * 0.72, page_height - 27, f"Level: {student_level}")

                all_detected = [item["qid"] for pg_items in question_positions.values()
                                for item in pg_items]
                if all_detected or mcq_results:
                    correct_count   = sum(1 for r in (mcq_results or []) if r.get("correct"))
                    incorrect_count = sum(1 for r in (mcq_results or []) if not r.get("correct"))
                    total_count     = len(all_detected) or len(mcq_results or [])

                    c.setFont("Helvetica-Bold", 11)
                    c.drawString(18, page_height - 46,
                                 f"Questions: {correct_count} correct  |  {incorrect_count} wrong/unattempted  (of {total_count})")

                    lx = page_width - 200
                    c.setFont("Helvetica", 9)
                    c.setFillColor(colors.Color(0.0, 0.65, 0.0))
                    c.drawString(lx,      page_height - 46, "\u2713 Correct")
                    c.setFillColor(colors.Color(0.85, 0.1, 0.1))
                    c.drawString(lx + 72, page_height - 46, "\u2717 Wrong / Unattempted")
            c.save()
            packet.seek(0)
            overlay_reader = PdfReader(packet)
            if overlay_reader.pages:
                page.merge_page(overlay_reader.pages[0])
            writer.add_page(page)

        output = BytesIO()
        writer.write(output)
        output.seek(0)
        return output.read()

    except Exception as e:
        print(f"[ERROR] Failed to create annotated PDF: {e}")
        return original_pdf_bytes

async def extract_text_from_upload(file: UploadFile) -> Dict[str, Any]:
    filename = getattr(file, "filename", "") or "upload"
    content_type = (getattr(file, "content_type", "") or "").lower()
    data = await file.read()

    if not data or len(data) < 20:
        return {"text": "", "kind": "unknown", "used_ocr": False, "needs_ocr": False, "error": "empty"}

    ext = (os.path.splitext(filename)[1] or "").lower()

    is_image = content_type.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}
    is_pdf = (content_type == "application/pdf") or ext == ".pdf"
    is_docx = (content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    }) or ext in {".docx", ".doc"}

    if is_image:
        try:
            return {"text": _clean_extracted_text(extract_text_from_image(data, filename=filename)),
                    "kind": "image", "used_ocr": True, "needs_ocr": False}
        except HTTPException as e:
            return {"text": "", "kind": "image", "used_ocr": True, "needs_ocr": False, "error": e.detail}

    if is_docx:
        try:
            return {"text": _clean_extracted_text(extract_text_from_docx(data, filename=filename)),
                    "kind": "docx", "used_ocr": False, "needs_ocr": False}
        except HTTPException as e:
            return {"text": "", "kind": "docx", "used_ocr": False, "needs_ocr": False, "error": e.detail}

    if is_pdf:
        info = extract_text_from_pdf(data, filename=filename)
        return {"text": info.get("text", ""), "kind": "pdf",
                "used_ocr": bool(info.get("used_ocr", False)),
                "needs_ocr": bool(info.get("needs_ocr", False)),
                "ocr_error": info.get("ocr_error")}

    # fallback: try as image
    try:
        return {"text": _clean_extracted_text(extract_text_from_image(data, filename=filename)),
                "kind": "unknown_as_image", "used_ocr": True, "needs_ocr": False}
    except Exception:
        return {"text": "", "kind": "unknown", "used_ocr": False, "needs_ocr": False,
                "error": f"Unsupported file type: {content_type or ext or 'unknown'}"}





@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/health/llm")
def health_llm():
    return {
        "ok": bool(gemini_client) and bool(GOOGLE_API_KEYS),
        "gemini": {
            "sdk_import_ok": genai is not None,
            "configured": bool(GOOGLE_API_KEYS),
            "num_keys_configured": len(GOOGLE_API_KEYS),
            "current_key_index": current_key_index + 1 if GOOGLE_API_KEYS else 0,
            "rate_limited_keys": list(rate_limited_keys),
            "client_ready": gemini_client is not None,
            "model": GEMINI_MODEL,
            "last_error": GEMINI_LAST_ERROR if GEMINI_LAST_ERROR else None,
        },
    }


@app.get("/homework/annotated-url/{homework_id}/{student_id}")
async def get_annotated_pdf_url(
    homework_id: int,
    student_id: int,
):
    """
    Get the URL for the annotated PDF.
    Returns JSON with the URL that can be used in your frontend.
    """
    base_url = get_public_base_url()
    return {
        "homework_id": homework_id,
        "student_id": student_id,
        "annotated_pdf_url": f"{base_url}/homework/annotated/{homework_id}/{student_id}"
    }
@app.get("/homework/annotated/{homework_id}/{student_id}")
async def get_annotated_pdf(
    homework_id: int,
    student_id: int,
):
    """
    Download the annotated PDF with tickmarks for a validated homework.
    This endpoint returns the PDF directly as a file download.
    """
    from fastapi.responses import Response
    
    try:
        # Fetch ERP record
        erp_row = fetch_student_record(homework_id, student_id)
        
        # Get submission image from ERP
        submission_image = erp_row.get("submission_image")
        if not submission_image:
            raise HTTPException(status_code=404, detail="No submission found")
        
        # Download the original file
        submission_url = STORAGE_BASE + submission_image
        resp = requests.get(submission_url, timeout=30)
        resp.raise_for_status()
        original_content = resp.content
        
        # Determine file type
        filename = submission_image.lower()
        is_pdf = filename.endswith('.pdf')
        
        if not is_pdf:
            raise HTTPException(status_code=400, detail="Annotated PDF only available for PDF submissions")
        
        # Get prompt and question type
        prompt = erp_row.get("prompt") or erp_row.get("question_prompt") or ""
        question_type = erp_row.get("question_type") or erp_row.get("type")
        student_level = fetch_student_level_from_erp(erp_row)
        
        # Extract text from PDF FIRST (needed for question type inference)
        pdf_info = extract_text_from_pdf(original_content, filename=submission_image)
        student_text = (pdf_info.get("text") or "").strip()
        
        if not student_text or len(student_text) < 10:
            raise HTTPException(status_code=400, detail="Could not extract text from PDF")
        
        final_question_type = (question_type or "").strip().lower()
        if final_question_type not in ("mcq", "narrative", "mixed"):
            final_question_type = infer_question_type_from_prompt(prompt, student_text)
        
        mcq_results = []
        status = "Needs Review"
        match_percentage = 0
        
        # Process based on question type
        if final_question_type == "mcq":
            correct = extract_correct_mcq_from_prompt(prompt)
            chosen = extract_mcq_choice(student_text)
            
            student_answers_by_qid = extract_mcq_answers_with_qid(student_text)
            
            if student_answers_by_qid:
                # Multiple MCQ
                parsed_questions = parse_questions_from_prompt(prompt)
                mcq_questions_with_answers = [q for q in parsed_questions if q.get('type') == 'mcq' and q.get('correct_answer')]
                
                for qid, student_ans in student_answers_by_qid.items():
                    matched = False
                    for pq in mcq_questions_with_answers:
                        pq_num = pq.get('qid', '').replace('Q', '').strip()
                        qid_num = qid.replace('Q', '').strip()
                        if pq_num == qid_num:
                            is_correct = student_ans.lower() == pq.get('correct_answer', '').lower()
                            mcq_results.append({
                                'qid': qid,
                                'chosen': student_ans,
                                'correct_answer': pq.get('correct_answer'),
                                'correct': is_correct,
                                'unattempted': False
                            })
                            matched = True
                            break
                    if not matched:
                        mcq_results.append({'qid': qid, 'chosen': student_ans, 'correct_answer': None, 'correct': False, 'unattempted': False})
                
                # Mark questions from the prompt that the student never answered
                answered_nums = {r['qid'].replace('Q', '').strip() for r in mcq_results}
                for pq in mcq_questions_with_answers:
                    pq_num = pq.get('qid', '').replace('Q', '').strip()
                    if pq_num not in answered_nums:
                        mcq_results.append({
                            'qid': pq.get('qid'),
                            'chosen': '',
                            'correct_answer': pq.get('correct_answer'),
                            'correct': False,
                            'unattempted': True
                        })
                
                if mcq_results:
                    correct_count = sum(1 for r in mcq_results if r.get('correct'))
                    mcq_credit = mcq_partial_credit(student_level)
                    match_percentage = int((correct_count * mcq_credit["credit_per_question"]) / max(1, len(mcq_results)))
                    status = "Verified" if match_percentage >= mcq_credit["passing_threshold"] else "Needs Review"
            elif correct and chosen:
                is_correct = (chosen == correct)
                mcq_credit = mcq_partial_credit(student_level)
                match_percentage = mcq_credit["credit_per_question"] if is_correct else 0
                status = "Verified" if match_percentage >= mcq_credit["passing_threshold"] else "Needs Review"
                _qid = extract_qid_from_prompt(prompt, erp_row)
                mcq_results = [{'qid': _qid, 'correct': is_correct, 'chosen': chosen, 'correct_answer': correct}]
        
        # For narrative, calculate score using AI
        if final_question_type == "narrative" and gemini_client:
            # Generate AI reference answer
            ai_prompt = (
                f"STUDENT_LEVEL: {student_level}\n"
                f"QUESTION:\n{prompt.strip()}\n\n"
                'Return ONLY valid JSON with keys: {"ai_reference_answer": string, "key_points": [string, ...]}.'
            )
            
            response_text = generate_gemini_response(
                prompt=ai_prompt,
                system_prompt="Generate a correct reference answer for homework evaluation. Keep it aligned with the student level. Output strict JSON only.",
                max_tokens=650,
                temperature=0.3,
            )
            
            if response_text:
                try:
                    import re
                    m = re.search(r'\{.*\}', response_text, flags=re.S)
                    payload = json.loads(m.group(0) if m else response_text)
                    
                    ai_reference_answer = (payload.get("ai_reference_answer") or "").strip()
                    key_points = payload.get("key_points") or []
                    
                    policy = level_policy(student_level)
                    sim = cosine_sim(student_text, ai_reference_answer)
                    covered, missing, coverage = keypoint_coverage(student_text, key_points, kp_threshold=policy["kp_thr"])
                    
                    final = policy["w_sim"] * sim + policy["w_cov"] * coverage
                    match_percentage = int(round(final * 100))
                    
                    if match_percentage >= policy["verified"]:
                        status = "Verified"
                    elif match_percentage >= policy["partial"]:
                        status = "Partial"
                    else:
                        status = "Needs Review"
                    
                    # Create result for narrative to show in PDF
                    if status == "Verified":
                        narrative_correct = True
                    elif status == "Partial":
                        narrative_correct = False
                    else:
                        narrative_correct = False
                    
                    _qid = extract_qid_from_prompt(prompt, erp_row)
                    mcq_results = [
                        {'qid': _qid, 'correct': narrative_correct, 'chosen': f'Score: {match_percentage}%', 'correct_answer': status}
                    ]
                except Exception as e:
                    print(f"[WARN] Failed to calculate narrative score: {e}")
        
        # Create annotated PDF
        annotated_pdf = create_annotated_pdf(
            original_pdf_bytes=original_content,
            mcq_results=mcq_results,
            match_percentage=match_percentage,
            status=status,
            student_level=student_level,
            question_type=final_question_type
        )
        
        # Return as file download
        return Response(
            content=annotated_pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"inline; filename=annotated_homework_{homework_id}_{student_id}.pdf"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Failed to generate annotated PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


def ai_evaluate_per_question(
    prompt: str,
    student_text: str,
    student_level: str = "Medium",
) -> list:
    """
    Uses Gemini AI to evaluate each question individually from the student's answer sheet.

    Logic:
      1. Send ALL questions from the prompt + the full student answer text to Gemini.
      2. Gemini identifies:
         - Which questions the student attempted (wrote any answer for)
         - For each attempted question: is the answer correct or wrong?
         - For unattempted questions: marks as unattempted
      3. Returns list of {qid, correct, unattempted, chosen, correct_answer}

    Marking rules (as per requirement):
      - No answer written            → unattempted  (orange ○)
      - Answer written, score < 35%  → wrong         (red ✗)
      - Answer written, score >= 35% → correct       (green ✓)
    """
    import re as _re

    # Parse question numbers from the homework prompt
    q_numbers = []
    for m in _re.finditer(r'\bQ(\d+)\s*[:\.]?', prompt, _re.IGNORECASE):
        n = int(m.group(1))
        if n not in q_numbers:
            q_numbers.append(n)
    q_numbers.sort()

    if not q_numbers:
        # Single question fallback — ask Gemini to evaluate it
        q_numbers = [1]

    total_q = len(q_numbers)
    qid_list = [f"Q{n}" for n in q_numbers]

    ai_prompt = f"""You are a strict homework evaluator. Your job is to evaluate each question individually.

STUDENT LEVEL: {student_level}
TOTAL QUESTIONS: {total_q}
QUESTION IDs: {', '.join(qid_list)}

HOMEWORK QUESTIONS (from teacher):
{prompt.strip()}

STUDENT'S ANSWER SHEET (OCR extracted):
{student_text.strip()}

INSTRUCTIONS:
For EACH question ({', '.join(qid_list)}), you must determine:
1. Did the student write ANY answer for this question? (check the student's answer sheet carefully)
2. If NO answer was written → mark as "unattempted"
3. If answer was written → evaluate if it is correct based on the homework question
   - score >= 35% similarity to correct answer → "correct"
   - score < 35% similarity → "wrong"

IMPORTANT RULES:
- Look carefully at student's answer text. The student may have written answers with question numbers like "1.", "Q1:", "1)", or just paragraph answers.
- If student's answer sheet is blank or has no relevant text for a specific question → unattempted
- Be strict: a vague or incomplete answer with < 35% match to expected answer = wrong
- A reasonably correct answer with >= 35% match = correct

Return ONLY valid JSON array, no extra text:
[
  {{"qid": "Q1", "status": "correct" | "wrong" | "unattempted", "student_answer_snippet": "brief snippet of what student wrote or empty string", "reason": "one line reason"}},
  ...one entry per question...
]"""

    response = generate_gemini_response(
        prompt=ai_prompt,
        system_prompt="You are a strict homework evaluator. Output only valid JSON array.",
        max_tokens=800,
        temperature=0.1,
    )

    results = []

    if response:
        try:
            # Extract JSON array from response
            m = re.search(r'\[.*\]', response, flags=re.S)
            if m:
                ai_data = json.loads(m.group(0))
                seen_qids = set()
                for item in ai_data:
                    qid = (item.get("qid") or "").strip()
                    status = (item.get("status") or "unattempted").strip().lower()
                    snippet = (item.get("student_answer_snippet") or "").strip()

                    if not qid:
                        continue
                    seen_qids.add(qid)

                    is_unattempted = status == "unattempted"
                    is_correct = status == "correct"

                    results.append({
                        "qid": qid,
                        "correct": is_correct,
                        "unattempted": is_unattempted,
                        "chosen": snippet[:80] if snippet else ("" if is_unattempted else "Answered"),
                        "correct_answer": "" if is_unattempted else ("Correct" if is_correct else "Wrong"),
                    })

                # Add any missing questions as unattempted
                for n in q_numbers:
                    qid = f"Q{n}"
                    if qid not in seen_qids:
                        results.append({
                            "qid": qid,
                            "correct": False,
                            "unattempted": True,
                            "chosen": "",
                            "correct_answer": "",
                        })

                # Sort by question number
                results.sort(key=lambda r: int(re.search(r'\d+', r.get("qid", "Q0")).group() or 0))
                return results

        except Exception as e:
            print(f"[WARN] ai_evaluate_per_question JSON parse failed: {e}, raw={response[:300]}")

    # ── Fallback: Gemini failed → use cosine similarity per question ──────────
    print("[WARN] ai_evaluate_per_question: Gemini failed, using cosine similarity fallback")
    return _cosine_fallback_per_question(prompt, student_text, q_numbers)


def _cosine_fallback_per_question(prompt: str, student_text: str, q_numbers: list) -> list:
    """
    Fallback when Gemini is unavailable.
    Extracts each question's answer from student text using regex segmentation,
    then scores with cosine similarity. < 35% = wrong, >= 35% = correct, no text = unattempted.
    """
    import re as _re

    WRONG_THRESHOLD = 0.35

    def q_start_regex(n):
        return _re.compile(
            rf'(?:q\s*{n}\s*[:.)\-]|(?<!\d){n}\s*[.):\-]\s|question\s*{n}\s*[:.)\-]?|ans(?:wer)?\s*{n}\s*[:.)\-]?)',
            _re.IGNORECASE
        )

    results = []
    text = student_text or ""

    for n in q_numbers:
        qid = f"Q{n}"
        pat = q_start_regex(n)
        m = pat.search(text)

        if not m:
            results.append({"qid": qid, "correct": False, "unattempted": True, "chosen": "", "correct_answer": ""})
            continue

        answer_start = m.end()
        answer_end = len(text)

        # Find where next question starts
        for other_n in q_numbers:
            if other_n == n:
                continue
            om = q_start_regex(other_n).search(text, answer_start)
            if om and om.start() < answer_end:
                answer_end = om.start()

        answer_text = text[answer_start:answer_end].strip()

        if not answer_text or len(answer_text.split()) < 2:
            results.append({"qid": qid, "correct": False, "unattempted": True, "chosen": "", "correct_answer": ""})
            continue

        # Score against full student text (rough proxy — no per-Q reference available here)
        sim = cosine_sim(answer_text, text)
        is_correct = sim >= WRONG_THRESHOLD

        results.append({
            "qid": qid,
            "correct": is_correct,
            "unattempted": False,
            "chosen": answer_text[:80],
            "correct_answer": "Correct" if is_correct else "Wrong",
        })

    return results


def build_per_question_results(
    prompt: str,
    student_text: str,
    overall_status: str,
    overall_score: int,
    ai_reference_answer: str = "",
    key_points: list = None,
    policy: dict = None,
    student_level: str = "Medium",
) -> list:
    """
    Main entry point for per-question evaluation.
    Delegates to ai_evaluate_per_question (Gemini) with cosine fallback.
    This replaces the old overall-status-based approach.
    """
    # Always use AI evaluation — it checks each question individually
    return ai_evaluate_per_question(prompt, student_text, student_level)


@app.post("/homework/validate")
async def homework_validate(
    student_id: int = Form(...),
    homework_id: int = Form(...),
    student_file: UploadFile = File(...),
):
    # 0) Fetch ERP record -> get all fields automatically
    erp_row = fetch_student_record(homework_id, student_id)
    
    # Extract fields from ERP record
    sub_institute_id = erp_row.get("sub_institute_id")
    syear = erp_row.get("syear")
    prompt = erp_row.get("prompt") or erp_row.get("question_prompt") or ""
    question_type = erp_row.get("question_type") or erp_row.get("type")
    
    student_level = fetch_student_level_from_erp(erp_row)
    policy = level_policy(student_level)

    # 2) Extract student text FIRST (needed for question type inference)
    student_info = await extract_text_from_upload(student_file)
    student_text = (student_info.get("text") or "").strip()
    
    # Keep a copy of the original file bytes for PDF annotation
    # Reset file cursor and read again
    await student_file.seek(0)
    original_file_bytes = await student_file.read()
    await student_file.seek(0)  # Reset for any further processing
    
    # Decide final question type: respect request value if valid, else infer using student text
    final_question_type = (question_type or "").strip().lower()
    if final_question_type not in ("mcq", "narrative", "mixed"):
        final_question_type = infer_question_type_from_prompt(prompt, student_text)

    # 1) Infer question_type from prompt automatically (NO EXTRA FIELD)
    # Try to parse mixed questions first
    parsed_questions = parse_questions_from_prompt(prompt)
    has_mcq = any(q.get('type') == 'mcq' for q in parsed_questions)
    has_narrative = any(q.get('type') == 'narrative' for q in parsed_questions)
    
    # Check if it's a PDF
    is_pdf_submission = student_info.get("kind") == "pdf"
    
    # Initialize annotated PDF filename
    annotated_pdf_filename = None
    annotated_pdf_url = None
    
    # Function to save annotated PDF — returns (filename, public_url)
    def save_annotated_pdf(pdf_bytes, hw_id, stud_id, results, score, stat, lvl, qtype="mcq"):
        if not pdf_bytes or len(pdf_bytes) < 100:
            return None, None
        try:
            outputs_dir = os.path.join(os.path.dirname(__file__), "outputs")
            os.makedirs(outputs_dir, exist_ok=True)
            ts = int(time.time())
            filename = f"marked_{hw_id}_{stud_id}_{ts}.pdf"
            filepath = os.path.join(outputs_dir, filename)
            
            annotated = create_annotated_pdf(
                original_pdf_bytes=pdf_bytes,
                mcq_results=results,
                match_percentage=score,
                status=stat,
                student_level=lvl,
                question_type=qtype
            )
            
            with open(filepath, "wb") as f:
                f.write(annotated)
            return filename, build_pdf_url(filename)
        except Exception as e:
            print(f"[WARN] Failed to save annotated PDF: {e}")
            return None, None

    MIN_WORDS = 3 if final_question_type == "mcq" else 8
    if len(student_text.split()) < MIN_WORDS:
        # Save annotated PDF even for unreadable (with status shown)
        if is_pdf_submission and original_file_bytes:
            # Show circle mark for unreadable
            unreadable_result = [{'qid': extract_qid_from_prompt(prompt, erp_row), 'correct': None, 'chosen': 'Unreadable', 'correct_answer': 'N/A'}]
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, unreadable_result, 0, "Unreadable", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": final_question_type,
            "student_level": student_level,
            "status": "Unreadable",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "Answer text could not be read clearly. Please upload a clearer file.",
            "student_extracted_text": student_text,
            "llm_used": False,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    if student_info.get("needs_ocr") and not student_text:
        # Save annotated PDF even for unreadable (with status shown)
        if is_pdf_submission and original_file_bytes:
            # Show circle mark for scanned PDF that needs OCR
            ocr_result = [{'qid': extract_qid_from_prompt(prompt, erp_row), 'correct': None, 'chosen': 'Needs OCR', 'correct_answer': 'N/A'}]
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, ocr_result, 0, "Unreadable", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": final_question_type,
            "student_level": student_level,
            "status": "Unreadable",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "This PDF looks scanned. OCR is required (install pdf2image + poppler) or upload a clearer file.",
            "student_extracted_text": student_text,
            "llm_used": False,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    
    if final_question_type == "mixed":
        # Process each question type separately and combine results
        mcq_results = []
        narrative_results = []
        
        # Extract ALL MCQ answers from student text with question numbers
        student_answers_by_qid = extract_mcq_answers_with_qid(student_text)
        
        # Extract MCQ answers from student text for each MCQ question
        for q in parsed_questions:
            if q.get('type') == 'mcq':
                qid = q.get('qid', '')
                q_num = qid.replace('Q', '').strip() if qid else ''
                
                # Try to get answer by question number first
                chosen = student_answers_by_qid.get(qid) or student_answers_by_qid.get(f"Q{q_num}")
                
                # Fallback to old method if no question number found
                if not chosen:
                    chosen = extract_mcq_choice(student_text)
                
                correct = q.get('correct_answer') or extract_correct_mcq_from_prompt(q.get('question', ''))
                
                if correct and chosen:
                    is_correct = (chosen.lower().strip() == correct.lower().strip())
                    mcq_results.append({
                        'qid': qid,
                        'correct': is_correct,
                        'chosen': chosen,
                        'correct_answer': correct,
                        'unattempted': False
                    })
                elif correct and not chosen:
                    # Student didn't answer this question at all
                    mcq_results.append({
                        'qid': qid,
                        'correct': False,
                        'chosen': '',
                        'correct_answer': correct,
                        'unattempted': True
                    })
        
        # For narrative questions, use AI to generate reference
        narrative_questions = [q for q in parsed_questions if q.get('type') == 'narrative']
        
        if narrative_questions and gemini_client:
            # Combine narrative questions into one prompt for AI
            narrative_prompt_text = "\n".join([
                f"{q.get('qid')}: {q.get('question')}" for q in narrative_questions
            ])
            
            ai_prompt = (
                f"STUDENT_LEVEL: {student_level}\n"
                f"QUESTIONS:\n{narrative_prompt_text}\n\n"
                'Return ONLY valid JSON with keys: {"ai_reference_answer": string, "key_points": [string, ...]}.'
            )
            
            response_text = generate_gemini_response(
                prompt=ai_prompt,
                system_prompt=(
                    "Generate correct reference answers for homework evaluation. "
                    "Keep it aligned with the student level. Output strict JSON only."
                ),
                max_tokens=650,
                temperature=0.3,
            )
            
            if response_text:
                try:
                    m = re.search(r'\{.*\}', response_text, flags=re.S)
                    payload = json.loads(m.group(0) if m else response_text)
                    
                    ai_reference_answer = (payload.get("ai_reference_answer") or "").strip()
                    key_points = payload.get("key_points") or []
                    
                    if isinstance(key_points, list):
                        key_points = [str(x).strip() for x in key_points if str(x).strip()]
                    
                    sim = cosine_sim(student_text, ai_reference_answer)
                    covered, missing, coverage = keypoint_coverage(
                        student_text, key_points, kp_threshold=policy["kp_thr"]
                    )
                    
                    final = policy["w_sim"] * sim + policy["w_cov"] * coverage
                    match_pct = int(round(final * 100))
                    
                    narrative_results = {
                        'similarity': sim,
                        'coverage': coverage,
                        'match_percentage': match_pct,
                        'key_points_covered': covered,
                        'key_points_missing': missing
                    }
                except Exception as e:
                    narrative_results = {'error': str(e)}
        
        # Calculate combined score with level-based partial credit for MCQ
        total_mcq = len(mcq_results)
        correct_mcq = sum(1 for r in mcq_results if r.get('correct'))
        
        # Get level-based credit per question
        mcq_credit = mcq_partial_credit(student_level)
        credit_per_q = mcq_credit["credit_per_question"]
        passing_threshold = mcq_credit["passing_threshold"]
        
        # Calculate MCQ score based on level (not just binary correct/incorrect)
        mcq_score = (correct_mcq * credit_per_q) / max(1, total_mcq)
        
        narrative_score = narrative_results.get('match_percentage', 0) if narrative_results else 0
        
        # Weight: 50% MCQ, 50% Narrative (if both exist)
        if total_mcq > 0 and narrative_results and 'error' not in narrative_results:
            final_score = int((mcq_score + narrative_score) / 2)
        elif total_mcq > 0:
            final_score = mcq_score
        elif narrative_results and 'error' not in narrative_results:
            final_score = narrative_score
        else:
            final_score = 0
        
        # Determine status
        if final_score >= policy["verified"]:
            status = "Verified"
        elif final_score >= policy["partial"]:
            status = "Partial"
        else:
            status = "Needs Review"
        
        # Save annotated PDF
        if is_pdf_submission and original_file_bytes and mcq_results:
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, mcq_results, final_score, status, student_level
            )
        
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": "mixed",
            "student_level": student_level,
            "status": status,
            "match_percentage": final_score,
            "submission_remarks": None,
            "rule_based_remark": f"MCQ: {correct_mcq}/{total_mcq} correct. Narrative score: {narrative_score}%. (Level: {student_level}, Credit per Q: {credit_per_q}%)",
            "llm_used": bool(narrative_results and 'error' not in narrative_results),
            "student_extracted_text": student_text,
            "mcq_results": mcq_results,
            "narrative_results": narrative_results,
            "question_marks": make_question_marks(mcq_results),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
            "debug": {
                "erp_row_fields": list(erp_row.keys()) if erp_row else [],
                "erp_student_level_raw": erp_row.get("student_level") or erp_row.get("level") or erp_row.get("difficulty") or erp_row.get("difficulty_level"),
                "mcq_credit_per_q": credit_per_q,
            },
        }

    elif final_question_type == "mcq":
        correct = extract_correct_mcq_from_prompt(prompt)
        chosen = extract_mcq_choice(student_text)
        
        # Try to extract multiple MCQ answers (for numbered questions like "1 A", "2 B")
        student_answers_by_qid = extract_mcq_answers_with_qid(student_text)
        has_multiple_mcq = len(student_answers_by_qid) > 1

        # Smart fallback: if answer looks like narrative (not MCQ), treat as narrative instead
        # This handles cases where question type is MCQ but student answered in narrative format
        # BUT if the answer contains Option A/B/C/D, treat as MCQ
        answer_has_mcq_option = bool(re.search(r"\b(option|answer|ans)\s*[:\-]?\s*[a-d]\b", _norm(student_text)))
        
        answer_looks_like_narrative = (
            len(student_text.split()) > 15 and  # More than 15 words
            not has_multiple_mcq and  # Not multiple numbered MCQ answers
            not answer_has_mcq_option  # No explicit option markers
        )

        # If answer looks like narrative, redirect to narrative processing
        if answer_looks_like_narrative and gemini_client:
            final_question_type = "narrative"
            redirect_to_narrative = True
        else:
            redirect_to_narrative = False
            
        # Handle multiple MCQ answers - grade each one
        if has_multiple_mcq:
            # Parse prompt for multiple correct answers
            parsed_questions = parse_questions_from_prompt(prompt)
            mcq_questions_with_answers = [q for q in parsed_questions if q.get('type') == 'mcq' and q.get('correct_answer')]
            
            # If we have correct answers for multiple questions, grade them
            if mcq_questions_with_answers:
                correct_count = 0
                total_count = len(student_answers_by_qid)
                mcq_results = []
                
                for qid, student_ans in student_answers_by_qid.items():
                    # Find matching correct answer
                    matched = False
                    for pq in mcq_questions_with_answers:
                        pq_num = pq.get('qid', '').replace('Q', '').strip()
                        qid_num = qid.replace('Q', '').strip()
                        if pq_num == qid_num:
                            is_correct = student_ans.lower() == pq.get('correct_answer', '').lower()
                            if is_correct:
                                correct_count += 1
                            mcq_results.append({
                                'qid': qid,
                                'chosen': student_ans,
                                'correct_answer': pq.get('correct_answer'),
                                'correct': is_correct,
                                'unattempted': False
                            })
                            matched = True
                            break
                    if not matched:
                        mcq_results.append({
                            'qid': qid,
                            'chosen': student_ans,
                            'correct_answer': None,
                            'correct': False,
                            'unattempted': False
                        })
                
                # Add any questions from the prompt that the student never answered
                answered_nums = {r['qid'].replace('Q', '').strip() for r in mcq_results}
                for pq in mcq_questions_with_answers:
                    pq_num = pq.get('qid', '').replace('Q', '').strip()
                    if pq_num not in answered_nums:
                        mcq_results.append({
                            'qid': pq.get('qid'),
                            'chosen': '',
                            'correct_answer': pq.get('correct_answer'),
                            'correct': False,
                            'unattempted': True
                        })
                
                # Calculate score based on level
                mcq_credit = mcq_partial_credit(student_level)
                credit_per_q = mcq_credit["credit_per_question"]
                match_percentage = int((correct_count * credit_per_q) / max(1, total_count))
                passing_threshold = mcq_credit["passing_threshold"]
                status = "Verified" if match_percentage >= passing_threshold else "Needs Review"
                
                # Save annotated PDF
                if is_pdf_submission and original_file_bytes:
                    annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                        original_file_bytes, homework_id, student_id, mcq_results, match_percentage, status, student_level
                    )
                
                return {
                    "student_id": student_id,
                    "homework_id": homework_id,
                    "sub_institute_id": sub_institute_id,
                    "syear": syear,
                    "question_type": "mcq",
                    "student_level": student_level,
                    "status": status,
                    "match_percentage": match_percentage,
                    "submission_remarks": None,
                    "rule_based_remark": f"Multiple MCQ: {correct_count}/{total_count} correct. Score: {match_percentage}% (Level: {student_level})",
                    "student_extracted_text": student_text,
                    "llm_used": False,
                    "question_marks": make_question_marks(mcq_results),
            "annotated_pdf": annotated_pdf_filename,
                    "debug": {"student_answers": student_answers_by_qid, "mcq_results": mcq_results},
                    "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
                }
            else:
                # No correct answers in prompt - return needs review with extracted answers
                # Save annotated PDF with circle mark
                if is_pdf_submission and original_file_bytes:
                    no_answer_result = [{'qid': extract_qid_from_prompt(prompt, erp_row), 'correct': None, 'chosen': 'No Answer Key', 'correct_answer': 'N/A'}]
                    annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                        original_file_bytes, homework_id, student_id, no_answer_result, 0, "Needs Review", student_level
                    )
                return {
                    "student_id": student_id,
                    "homework_id": homework_id,
                    "sub_institute_id": sub_institute_id,
                    "syear": syear,
                    "question_type": "mcq",
                    "student_level": student_level,
                    "status": "Needs Review",
                    "match_percentage": 0,
                    "submission_remarks": None,
                    "rule_based_remark": f"Found {len(student_answers_by_qid)} MCQ answers but no correct answers in prompt. Include 'Correct: B' for each question.",
                    "student_extracted_text": student_text,
                    "llm_used": False,
                    "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
                    "debug": {"student_answers": student_answers_by_qid, "correct_answers_in_prompt": False},
                    "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
                }

        if redirect_to_narrative:
            pass  # Will continue to narrative handling
        elif not correct:
            # Save annotated PDF with circle mark
            if is_pdf_submission and original_file_bytes:
                no_correct_result = [{'qid': extract_qid_from_prompt(prompt, erp_row), 'correct': None, 'chosen': 'Not Found', 'correct_answer': 'N/A'}]
                annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                    original_file_bytes, homework_id, student_id, no_correct_result, 0, "Needs Review", student_level
                )
            return {
                "student_id": student_id,
                "homework_id": homework_id,
                "sub_institute_id": sub_institute_id,
                "syear": syear,
                "question_type": "mcq",
                "student_level": student_level,
                "status": "Needs Review",
                "match_percentage": 0,
                "submission_remarks": None,
                "rule_based_remark": "MCQ correct option not found in prompt. Include 'Correct: B' or similar in prompt.",
                "student_extracted_text": student_text,
                "llm_used": False,
                "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
                "debug": {"correct": correct, "chosen": chosen},
                "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
            }
        elif not chosen:
            # Save annotated PDF with circle mark
            if is_pdf_submission and original_file_bytes:
                no_chosen_result = [{'qid': extract_qid_from_prompt(prompt, erp_row), 'correct': None, 'chosen': 'Not Detected', 'correct_answer': correct or 'N/A'}]
                annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                    original_file_bytes, homework_id, student_id, no_chosen_result, 0, "Needs Review", student_level
                )
            return {
                "student_id": student_id,
                "homework_id": homework_id,
                "sub_institute_id": sub_institute_id,
                "syear": syear,
                "question_type": "mcq",
                "student_level": student_level,
                "status": "Needs Review",
                "match_percentage": 0,
                "submission_remarks": None,
                "rule_based_remark": "Student option (A/B/C/D) not detected clearly.",
                "student_extracted_text": student_text,
                "llm_used": False,
                "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
                "debug": {"correct": correct, "chosen": chosen},
                "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
            }

        # Only process MCQ validation if not redirecting to narrative
        if not redirect_to_narrative:
            is_correct = (chosen == correct)
            
            # Get level-based credit
            mcq_credit = mcq_partial_credit(student_level)
            credit_per_q = mcq_credit["credit_per_question"]
            
            # Calculate score based on level
            match_percentage = credit_per_q if is_correct else 0
            
            # Determine status based on level threshold
            passing_threshold = mcq_credit["passing_threshold"]
            status = "Verified" if match_percentage >= passing_threshold else "Needs Review"
            
            # Save annotated PDF
            _qid = extract_qid_from_prompt(prompt, erp_row)
            mcq_results_single = [{'qid': _qid, 'correct': is_correct, 'chosen': chosen, 'correct_answer': correct}]
            if is_pdf_submission and original_file_bytes:
                annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                    original_file_bytes, homework_id, student_id, mcq_results_single, match_percentage, status, student_level
                )
            
            return {
                "student_id": student_id,
                "homework_id": homework_id,
                "sub_institute_id": sub_institute_id,
                "syear": syear,
                "question_type": "mcq",
                "student_level": student_level,
                "status": status,
                "match_percentage": match_percentage,
                "submission_remarks": None,
                "rule_based_remark": f"{'Correct' if is_correct else 'Incorrect'}. Score: {match_percentage}% (Level: {student_level}, Credit per Q: {credit_per_q}%)",
                "student_extracted_text": student_text,
                "llm_used": False,
                "question_marks": make_question_marks(mcq_results_single),
            "annotated_pdf": annotated_pdf_filename,
                "debug": {"correct": correct, "chosen": chosen, "level": student_level, "credit_per_q": credit_per_q},
                "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
            }

 
    if gemini_client is None:
        # Save annotated PDF
        if is_pdf_submission and original_file_bytes:
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, [], 0, "Needs Review", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": "narrative",
            "student_level": student_level,
            "status": "Needs Review",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "Gemini not configured. Check /health/llm.",
            "llm_used": False,
            "llm_error": parse_gemini_error(GEMINI_LAST_ERROR),
            "student_extracted_text": student_text,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    user_prompt = (
        f"STUDENT_LEVEL: {student_level}\n"
        f"QUESTION:\n{prompt.strip()}\n\n"
        'Return ONLY valid JSON with keys: {"ai_reference_answer": string, "key_points": [string, ...]}.'
    )

    response_text = generate_gemini_response(
        prompt=user_prompt,
        system_prompt=(
            "Generate a correct reference answer for homework evaluation. "
            "Keep it aligned with the student level. Output strict JSON only."
        ),
        max_tokens=650,
        temperature=0.3,
    )

    if not response_text:
        # Save annotated PDF
        if is_pdf_submission and original_file_bytes:
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, [], 0, "Needs Review", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": "narrative",
            "student_level": student_level,
            "status": "Needs Review",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "Gemini failed. Check /health/llm.",
            "llm_used": False,
            "llm_error": parse_gemini_error(GEMINI_LAST_ERROR),
            "student_extracted_text": student_text,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    try:
        m = re.search(r"\{.*\}", response_text, flags=re.S)
        payload = json.loads(m.group(0) if m else response_text)
    except Exception as e:
        # Save annotated PDF
        if is_pdf_submission and original_file_bytes:
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, [], 0, "Needs Review", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": "narrative",
            "student_level": student_level,
            "status": "Needs Review",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "Gemini returned non-JSON output.",
            "llm_used": False,
            "llm_error": {"ok": False, "error_type": "GEMINI_BAD_JSON", "message": str(e), "raw": response_text[:800]},
            "student_extracted_text": student_text,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    ai_reference_answer = (payload.get("ai_reference_answer") or "").strip()
    key_points = payload.get("key_points") or []
    if not isinstance(key_points, list):
        key_points = []
    key_points = [str(x).strip() for x in key_points if str(x).strip()]

    if not ai_reference_answer:
        # Save annotated PDF
        if is_pdf_submission and original_file_bytes:
            annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
                original_file_bytes, homework_id, student_id, [], 0, "Needs Review", student_level
            )
        return {
            "student_id": student_id,
            "homework_id": homework_id,
            "sub_institute_id": sub_institute_id,
            "syear": syear,
            "question_type": "narrative",
            "student_level": student_level,
            "status": "Needs Review",
            "match_percentage": 0,
            "submission_remarks": None,
            "rule_based_remark": "AI returned empty reference answer.",
            "llm_used": True,
            "student_extracted_text": student_text,
            "question_marks": make_question_marks([]),
            "annotated_pdf": annotated_pdf_filename,
            "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
        }

    sim = cosine_sim(student_text, ai_reference_answer)
    covered, missing, coverage = keypoint_coverage(student_text, key_points, kp_threshold=policy["kp_thr"])

    final = policy["w_sim"] * sim + policy["w_cov"] * coverage
    match_pct = int(round(final * 100))

    if match_pct >= policy["verified"]:
        status = "Verified"
    elif match_pct >= policy["partial"]:
        status = "Partial"
    else:
        status = "Needs Review"

    # Short remark (Gemini), fallback to rule-based
    remark_prompt = (
        f"Student level: {student_level}\n"
        f"Match: {match_pct}%\n"
        f"Missing key points: {missing[:6]}\n\n"
        "Write a short, factual teacher remark (2-4 lines). No marks. No overpraise."
    )

    resp2_prompt = (
        f"REFERENCE ANSWER:\n{ai_reference_answer[:900]}\n\n"
        f"STUDENT ANSWER:\n{student_text[:900]}\n\n"
        f"{remark_prompt}"
    )

    submission_remark = generate_gemini_response(
        prompt=resp2_prompt,
        system_prompt="You are a strict, helpful teacher. Be concise and factual.",
        max_tokens=140,
        temperature=0.6,
    )

    rule_based_remark = None
    remark_llm_used = bool(submission_remark)
    remark_llm_error = None if submission_remark else (GEMINI_LAST_ERROR or "Unknown LLM error")

    if not submission_remark:
        if status == "Verified":
            rule_based_remark = "Homework matches the expected answer well. Good coverage of the key ideas."
        elif status == "Partial":
            rule_based_remark = "Homework is partially correct. Improve coverage of missing key points and make the explanation clearer."
        else:
            rule_based_remark = "Homework does not match the expected answer enough. Please review the topic and resubmit with clearer, complete points."

    # Save annotated PDF — evaluate EACH question individually against student text
    per_question_results = build_per_question_results(
        prompt, student_text, status, match_pct,
        ai_reference_answer=ai_reference_answer,
        key_points=key_points,
        policy=policy,
        student_level=student_level,
    )
    if is_pdf_submission and original_file_bytes:
        annotated_pdf_filename, annotated_pdf_url = save_annotated_pdf(
            original_file_bytes, homework_id, student_id, per_question_results, match_pct, status, student_level, "narrative"
        )

    return {
        "student_id": student_id,
        "homework_id": homework_id,
        "sub_institute_id": sub_institute_id,
        "syear": syear,
        "question_type": "narrative",
        "student_level": student_level,
        "status": status,
        "match_percentage": match_pct,
        "submission_remarks": submission_remark if submission_remark else None,
        "rule_based_remark": rule_based_remark,
        "llm_used": True,
        "remark_llm_used": remark_llm_used,
        "remark_llm_error": remark_llm_error,
        "student_extracted_text": student_text,
        "ai_reference_answer": ai_reference_answer,
        "key_points": key_points,
        "key_points_covered": covered,
        "key_points_missing": missing,
        "question_marks": make_question_marks(per_question_results),
        "annotated_pdf": annotated_pdf_filename,
        "debug": {
            "similarity": sim,
            "coverage": coverage,
            "policy": policy,
            "per_question_results": per_question_results,
            "erp_row_fields": list(erp_row.keys()) if erp_row else [],
            "erp_student_level_raw": erp_row.get("student_level") or erp_row.get("level") or erp_row.get("difficulty") or erp_row.get("difficulty_level"),
        },
        "extraction": {"student": {k: v for k, v in student_info.items() if k != "text"}},
    }